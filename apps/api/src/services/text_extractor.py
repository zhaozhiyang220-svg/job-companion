import io
from pathlib import Path

import pypdfium2
from docx import Document


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _pdf(data)
    if ext in {".docx", ".doc"}:
        return _docx(data)
    raise ValueError(f"unsupported file extension: {ext}")


def _pdf(data: bytes) -> str:
    pdf = pypdfium2.PdfDocument(data)
    chunks: list[str] = []
    for page in pdf:
        text_page = page.get_textpage()
        chunks.append(text_page.get_text_range() or "")
    return "\n".join(chunks)


def _docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
